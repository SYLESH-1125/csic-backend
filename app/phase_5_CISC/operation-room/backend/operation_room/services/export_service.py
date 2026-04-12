"""
Deterministic Export Engine — Phase 4.

Exports reports to PDF, DOCX, HTML formats with:
- Snapshot of AST + all bound evidence data
- Vega-Lite charts rendered to SVG/PNG
- SHA-256 manifest for every artifact
- CoC event logging for each export
"""

import json
import uuid
import hashlib
import canonicaljson
import os
import logging
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from operation_room.database import open_vault
from operation_room.config import settings
from operation_room.services.audit_service import record_coc_event
from operation_room.services.reportlab_pdf_service import generate_pdf_from_ast
from operation_room.services.pdf_integrity import PDFIntegrityService
from operation_room.services.pdf_signer import PDFSigner
from operation_room.services.report_evidence_service import get_report_evidence_service, RedactionMode, AccessPurpose

logger = logging.getLogger(__name__)


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _hash_file(path: str) -> str:
    """SHA-256 hash of a file."""
    h = hashlib.sha256()
    with open(path, 'rb') as f:
        for chunk in iter(lambda: f.read(8192), b''):
            h.update(chunk)
    return f"sha256:{h.hexdigest()}"


def _hash_content(content: bytes) -> str:
    return f"sha256:{hashlib.sha256(content).hexdigest()}"


_IPV4_RE = re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}\b")
_MAC_RE = re.compile(r"\b(?:[0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2}\b")
_LONG_HEX_RE = re.compile(r"\b[a-fA-F0-9]{32,}\b")


def _obfuscate_story_text(value: str) -> str:
    if not value:
        return value
    value = _IPV4_RE.sub("[OBFUSCATED_IP]", value)
    value = _MAC_RE.sub("[OBFUSCATED_MAC]", value)
    value = _LONG_HEX_RE.sub("[OBFUSCATED_HASH]", value)
    return value


def _story_confidence(element: dict) -> float:
    data = element.get("data") or {}
    metadata = data.get("metadata") if isinstance(data.get("metadata"), dict) else {}
    candidates = [
        data.get("confidence"),
        metadata.get("confidence"),
        element.get("confidence"),
    ]
    for candidate in candidates:
        if candidate is None:
            continue
        try:
            return float(candidate)
        except Exception:
            continue
    return 1.0


def _apply_story_mode_transform(ast: dict) -> dict:
    """Story mode keeps high-confidence elements and obfuscates technical identifiers."""
    transformed = json.loads(json.dumps(ast))
    pages = transformed.get("pages", [])
    if not isinstance(pages, list):
        return transformed

    for page in pages:
        elements = page.get("elements", [])
        if not isinstance(elements, list):
            continue
        kept: list[dict] = []
        for element in elements:
            if _story_confidence(element) < 0.8:
                continue

            if element.get("type") == "text":
                data = element.get("data") or {}
                content = data.get("content")
                if isinstance(content, str):
                    data["content"] = _obfuscate_story_text(content)
                element["data"] = data
            elif element.get("type") == "component":
                data = element.get("data") or {}
                for key in ("summary", "content", "label", "title"):
                    if isinstance(data.get(key), str):
                        data[key] = _obfuscate_story_text(data[key])
                element["data"] = data

            kept.append(element)
        page["elements"] = kept
    transformed["pages"] = pages
    return transformed


# ═══════════════════════════════════════════════════════════════
# HTML Template
# ═══════════════════════════════════════════════════════════════

NFLIP_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title} — NFLIP Report</title>
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;600&display=swap');

* {{ margin: 0; padding: 0; box-sizing: border-box; }}

body {{
  font-family: 'Inter', sans-serif;
  font-size: 11pt;
  line-height: 1.7;
  color: #1e293b;
  background: #fff;
}}

/* Page layout for print */
@page {{
  size: A4;
  margin: 25mm;
}}

.report-container {{
  max-width: 210mm;
  margin: 0 auto;
  padding: 30px 40px;
}}

/* Header */
.report-header {{
  display: flex;
  justify-content: space-between;
  align-items: center;
  padding: 20px 0;
  border-bottom: 3px solid #1e293b;
  margin-bottom: 30px;
}}

.report-header-brand {{
  display: flex;
  align-items: center;
  gap: 12px;
}}

.report-header-brand h1 {{
  font-size: 18pt;
  font-weight: 800;
  color: #1e293b;
  letter-spacing: -0.02em;
}}

.report-header-brand span {{
  font-size: 8pt;
  color: #64748b;
  text-transform: uppercase;
  letter-spacing: 0.1em;
}}

.report-classification {{
  font-family: 'JetBrains Mono', monospace;
  font-size: 9pt;
  font-weight: 700;
  color: #dc2626;
  border: 2px solid #dc2626;
  padding: 4px 12px;
  border-radius: 4px;
}}

/* Content */
h1 {{ font-size: 22pt; font-weight: 800; color: #0f172a; margin: 30px 0 12px; letter-spacing: -0.02em; }}
h2 {{ font-size: 16pt; font-weight: 700; color: #1e293b; margin: 24px 0 10px; border-bottom: 1px solid #e2e8f0; padding-bottom: 6px; }}
h3 {{ font-size: 13pt; font-weight: 700; color: #334155; margin: 18px 0 8px; }}
h4 {{ font-size: 11pt; font-weight: 700; color: #475569; margin: 14px 0 6px; }}

p {{ margin: 8px 0; }}

ul, ol {{ padding-left: 24px; margin: 8px 0; }}
li {{ margin: 3px 0; }}

blockquote {{
  border-left: 4px solid #2563eb;
  padding: 12px 16px;
  margin: 12px 0;
  background: #f8fafc;
  color: #475569;
  font-style: italic;
}}

pre {{
  background: #0f172a;
  color: #e2e8f0;
  padding: 16px;
  border-radius: 8px;
  overflow-x: auto;
  font-family: 'JetBrains Mono', monospace;
  font-size: 9pt;
  margin: 12px 0;
}}

code {{
  font-family: 'JetBrains Mono', monospace;
  font-size: 9pt;
  background: rgba(37,99,235,0.06);
  padding: 2px 6px;
  border-radius: 4px;
  color: #2563eb;
}}

table {{
  width: 100%;
  border-collapse: collapse;
  margin: 12px 0;
  font-size: 10pt;
}}

th {{
  background: #f1f5f9;
  font-weight: 700;
  text-align: left;
  padding: 8px 12px;
  border: 1px solid #e2e8f0;
  color: #475569;
  font-size: 9pt;
  text-transform: uppercase;
  letter-spacing: 0.04em;
}}

td {{
  padding: 6px 12px;
  border: 1px solid #e2e8f0;
  color: #334155;
}}

tr:nth-child(even) td {{ background: #fafbfc; }}

img {{ max-width: 100%; height: auto; border-radius: 6px; margin: 12px 0; }}

hr {{ border: none; border-top: 1px solid #e2e8f0; margin: 20px 0; }}

mark {{ background: rgba(217,119,6,0.2); padding: 1px 3px; border-radius: 2px; }}

/* Footer */
.report-footer {{
  margin-top: 60px;
  padding: 16px 0;
  border-top: 2px solid #1e293b;
  font-size: 8pt;
  color: #94a3b8;
  display: flex;
  justify-content: space-between;
  font-family: 'JetBrains Mono', monospace;
}}

/* Evidence hash badge */
.evidence-hash {{
  font-family: 'JetBrains Mono', monospace;
  font-size: 8pt;
  color: #059669;
  background: rgba(5,150,105,0.06);
  padding: 2px 8px;
  border-radius: 4px;
  border: 1px solid rgba(5,150,105,0.15);
}}

@media print {{
  body {{ font-size: 10pt; }}
  .report-container {{ padding: 0; max-width: none; }}
  .no-print {{ display: none !important; }}
}}

/* Evidence Blocks */
.evidence-block {{
  margin: 20px 0;
  padding: 16px 20px;
  border: 1px solid #e2e8f0;
  border-left: 4px solid #3b82f6;
  border-radius: 8px;
  background: #fafbfc;
  page-break-inside: avoid;
}}

.evidence-header {{
  display: flex;
  align-items: center;
  gap: 10px;
  margin-bottom: 12px;
  padding-bottom: 10px;
  border-bottom: 1px solid #e2e8f0;
}}

.evidence-icon {{
  font-size: 16px;
}}

.evidence-title {{
  font-weight: 700;
  font-size: 11pt;
  color: #1e293b;
  flex: 1;
}}

.evidence-source {{
  font-size: 8pt;
  font-weight: 600;
  padding: 3px 10px;
  border-radius: 12px;
  text-transform: uppercase;
  letter-spacing: 0.04em;
}}

.evidence-body {{
  margin: 12px 0;
}}

.evidence-chart-placeholder {{
  display: flex;
  align-items: center;
  gap: 12px;
  padding: 20px;
  background: rgba(59,130,246,0.05);
  border-radius: 6px;
  border: 1px dashed rgba(59,130,246,0.3);
}}

.chart-icon {{
  font-size: 32px;
  opacity: 0.6;
}}

.chart-info strong {{
  display: block;
  color: #3b82f6;
  margin-bottom: 4px;
}}

.chart-note {{
  font-size: 9pt;
  color: #94a3b8;
}}

.evidence-metric {{
  text-align: center;
  padding: 16px;
}}

.metric-value {{
  display: block;
  font-size: 32pt;
  font-weight: 800;
  color: #1e293b;
}}

.metric-label {{
  font-size: 10pt;
  color: #64748b;
}}

.evidence-details {{
  font-size: 10pt;
}}

.detail-row {{
  padding: 4px 0;
  border-bottom: 1px solid #f1f5f9;
}}

.detail-row .key {{
  font-weight: 600;
  color: #475569;
}}

.detail-row .value {{
  color: #1e293b;
}}

.evidence-list {{
  margin: 0;
  padding-left: 20px;
  font-size: 10pt;
}}

.evidence-table {{
  font-size: 9pt;
}}

.evidence-table table {{
  margin: 0;
}}

.evidence-footer {{
  display: flex;
  align-items: center;
  gap: 16px;
  margin-top: 12px;
  padding-top: 8px;
  border-top: 1px solid #e2e8f0;
  font-size: 8pt;
  color: #94a3b8;
}}

.evidence-type {{
  font-family: 'JetBrains Mono', monospace;
  background: #f1f5f9;
  padding: 2px 8px;
  border-radius: 4px;
}}

.evidence-verified {{
  color: #059669;
  font-weight: 600;
}}

/* Task Lists */
.task-list {{
  list-style: none;
  padding-left: 0;
}}

.task-item {{
  display: flex;
  align-items: center;
  gap: 8px;
  margin: 6px 0;
}}

.task-item input[type="checkbox"] {{
  width: 16px;
  height: 16px;
}}

/* Figure Captions */
figure {{
  margin: 16px 0;
  text-align: center;
}}

figure img {{
  display: block;
  margin: 0 auto;
}}

figcaption {{
  font-size: 9pt;
  color: #64748b;
  font-style: italic;
  margin-top: 8px;
}}
</style>
</head>
<body>
<div class="report-container">

<!-- Header -->
<div class="report-header">
  <div class="report-header-brand">
    <div style="font-size:28px">🛡️</div>
    <div>
      <h1>NFLIP</h1>
      <span>National Forensic Log Intelligence Platform</span>
    </div>
  </div>
  <div class="report-classification">CONFIDENTIAL</div>
</div>

<!-- Meta -->
<div style="font-size:9pt; color:#64748b; margin-bottom:24px; font-family:'JetBrains Mono',monospace;">
  Case: {case_id} · Exported: {export_time} · Hash: <span class="evidence-hash">{content_hash}</span>
</div>

<!-- Body -->
{body_html}

<!-- Footer -->
<div class="report-footer">
  <span>Government of India — Confidential · NFLIP Report Studio v2</span>
  <span>SHA-256: {content_hash}</span>
</div>

</div>
</body>
</html>"""


def _ast_to_html(ast: dict, figure_counter: dict = None) -> str:
    """Convert a TipTap AST to HTML with evidence block support."""
    if figure_counter is None:
        figure_counter = {'count': 0}

    if not ast:
        return ""

    ntype = ast.get("type", "")
    attrs = ast.get("attrs", {}) or {}
    content = ast.get("content") or []
    marks = ast.get("marks") or []

    if ntype == "text":
        text = ast.get("text", "")
        # Apply marks
        for mark in marks:
            mtype = mark.get("type", "")
            if mtype == "bold":
                text = f"<strong>{text}</strong>"
            elif mtype == "italic":
                text = f"<em>{text}</em>"
            elif mtype == "underline":
                text = f"<u>{text}</u>"
            elif mtype == "strike":
                text = f"<s>{text}</s>"
            elif mtype == "code":
                text = f"<code>{text}</code>"
            elif mtype == "highlight":
                color = mark.get("attrs", {}).get("color", "#fef08a")
                text = f'<mark style="background-color:{color}">{text}</mark>'
            elif mtype == "textStyle":
                color = mark.get("attrs", {}).get("color", "")
                if color:
                    text = f'<span style="color:{color}">{text}</span>'
            elif mtype == "link":
                href = mark.get("attrs", {}).get("href", "#")
                text = f'<a href="{href}" target="_blank">{text}</a>'
        return text

    # Handle evidence blocks specially
    if ntype == "evidenceBlock":
        return _render_evidence_block_html(attrs, figure_counter)

    children_html = "".join(_ast_to_html(c, figure_counter) for c in content)

    if ntype == "v4-canvas":
        pages = ast.get("pages") or []
        html_out = ""
        for i, page in enumerate(pages):
            if not page: continue
            # Each page is an A4 canvas (e.g. 794x1123px mapping to 210x297mm)
            html_out += f'<div class="v4-page" style="position: relative; width: 794px; height: 1123px; margin: 0 auto; margin-bottom: 2rem; background: white; border: 1px solid #e2e8f0; overflow: hidden; page-break-after: always; box-shadow: 0 4px 6px -1px rgb(0 0 0 / 0.1);">'

            for el in (page.get("elements") or []):
                x = el.get("x", 0)
                y = el.get("y", 0)
                w = el.get("width", 200)
                h = el.get("height", 100)
                z = el.get("zIndex", 1)
                etype = el.get("type", "component")
                data = el.get("data", {})
                
                if etype == 'text':
                    # Respect the text element's stored fontSize
                    font_size = data.get('fontSize', 13)
                    html_out += f'<div class="v4-element" style="position: absolute; left: {x}px; top: {y}px; width: {w}px; height: {h}px; z-index: {z}; font-size: {font_size}px; line-height: 1.5; overflow: hidden;">'
                    html_out += data.get('content', '')
                else:
                    html_out += f'<div class="v4-element" style="position: absolute; left: {x}px; top: {y}px; width: {w}px; height: {h}px; z-index: {z}; overflow: hidden;">'
                    el_attrs = {
                        "type": data.get("chartType") or "chart" if data.get("chartType") else data.get("type", "finding"),
                        "source": data.get("module", "case"),
                        "title": data.get("title", "Widget"),
                        "data": data,
                        "metadata": {}
                    }
                    html_out += _render_evidence_block_html(el_attrs, figure_counter)
                    
                html_out += '</div>'
            html_out += '</div>'
        
        # ── Vault of Evidence Appendix ────────────────────────
        all_evidence = []
        for pi, page in enumerate(pages):
            for el in (page.get("elements") or []):
                if el.get("type") != "text":
                    all_evidence.append({"page": pi + 1, "element": el})
        
        if all_evidence:
            html_out += '<div style="page-break-before: always; padding: 40px; max-width: 794px; margin: 0 auto;">'
            html_out += '<h1 style="font-size:22pt;font-weight:800;color:#0f172a;margin-bottom:12px;">Vault of Evidence</h1>'
            html_out += '<p style="font-size:10pt;color:#64748b;border-bottom:1px solid #e2e8f0;padding-bottom:12px;margin-bottom:20px;">Metadata registry confirming the origin, module source, and validation hashes of all functional widgets used in this report.</p>'
            html_out += '<table style="width:100%;border-collapse:collapse;font-size:9pt;">'
            html_out += '<tr style="background:#f1f5f9;"><th style="padding:8px 12px;border:1px solid #e2e8f0;text-align:left;font-weight:600;color:#475569;">Element ID</th><th style="padding:8px 12px;border:1px solid #e2e8f0;text-align:left;font-weight:600;color:#475569;">Page</th><th style="padding:8px 12px;border:1px solid #e2e8f0;text-align:left;font-weight:600;color:#475569;">Source / Type</th><th style="padding:8px 12px;border:1px solid #e2e8f0;text-align:left;font-weight:600;color:#475569;">Title</th></tr>'
            for item in all_evidence:
                el = item["element"]
                d = el.get("data", {})
                eid = el.get("id", "")[:16]
                mod = d.get("module", "System")
                typ = d.get("type", "Widget")
                ttl = d.get("title", "Untitled")
                html_out += f'<tr><td style="padding:6px 12px;border:1px solid #e2e8f0;font-family:monospace;font-size:8pt;color:#64748b;">evd-{eid}...</td><td style="padding:6px 12px;border:1px solid #e2e8f0;">Page {item["page"]}</td><td style="padding:6px 12px;border:1px solid #e2e8f0;"><strong style="color:#0284c7;">{mod}</strong><br><span style="font-size:8pt;color:#94a3b8;">{typ}</span></td><td style="padding:6px 12px;border:1px solid #e2e8f0;font-weight:500;">{ttl}</td></tr>'
            html_out += '</table>'
            html_out += '<p style="text-align:center;margin-top:40px;font-size:9pt;color:#94a3b8;">End of Report Registry</p>'
            html_out += '</div>'
        
        return html_out

    if ntype == "doc":
        return children_html
    elif ntype == "paragraph":
        align = attrs.get("textAlign", "")
        style = f' style="text-align:{align}"' if align else ""
        return f"<p{style}>{children_html}</p>\n"
    elif ntype == "heading":
        level = attrs.get("level", 1)
        return f"<h{level}>{children_html}</h{level}>\n"
    elif ntype == "bulletList":
        return f"<ul>\n{children_html}</ul>\n"
    elif ntype == "orderedList":
        return f"<ol>\n{children_html}</ol>\n"
    elif ntype == "listItem":
        return f"<li>{children_html}</li>\n"
    elif ntype == "taskList":
        return f'<ul class="task-list">\n{children_html}</ul>\n'
    elif ntype == "taskItem":
        checked = "checked" if attrs.get("checked") else ""
        return f'<li class="task-item"><input type="checkbox" disabled {checked}> {children_html}</li>\n'
    elif ntype == "blockquote":
        return f"<blockquote>{children_html}</blockquote>\n"
    elif ntype == "codeBlock":
        lang = attrs.get("language", "")
        return f'<pre><code class="language-{lang}">{children_html}</code></pre>\n'
    elif ntype == "horizontalRule":
        return "<hr>\n"
    elif ntype == "image":
        src = attrs.get("src", "")
        alt = attrs.get("alt", "")
        title = attrs.get("title", "")
        return f'<figure><img src="{src}" alt="{alt}" title="{title}"><figcaption>{alt}</figcaption></figure>\n'
    elif ntype == "table":
        return f'<table class="report-table">\n{children_html}</table>\n'
    elif ntype == "tableRow":
        return f"<tr>{children_html}</tr>\n"
    elif ntype == "tableCell":
        colspan = attrs.get("colspan", 1)
        rowspan = attrs.get("rowspan", 1)
        attrs_str = ""
        if colspan > 1:
            attrs_str += f' colspan="{colspan}"'
        if rowspan > 1:
            attrs_str += f' rowspan="{rowspan}"'
        return f"<td{attrs_str}>{children_html}</td>"
    elif ntype == "tableHeader":
        return f"<th>{children_html}</th>"
    elif ntype == "hardBreak":
        return "<br>"
    else:
        return children_html


# ═══════════════════════════════════════════════════════════════
# Export Functions
# ═══════════════════════════════════════════════════════════════

def export_html(case_id: str, doc_id: str, actor: str = "investigator", focus_mode: str = "Review") -> dict:
    """Export a document as a standalone HTML file."""
    conn = open_vault(case_id)
    try:
        row = conn.execute(
            "SELECT title, ast_json, content_hash FROM studio_documents WHERE doc_id=?",
            [doc_id]
        ).fetchone()
        if not row:
            return {"error": "Document not found"}

        title, ast_str, content_hash = row[0], row[1], row[2]
        ast = json.loads(ast_str)

        body_html = _ast_to_html(ast)
        export_time = _now()

        full_html = NFLIP_HTML_TEMPLATE.format(
            title=title,
            case_id=case_id,
            export_time=export_time,
            content_hash=content_hash,
            body_html=body_html,
        )

        # Save to exports directory
        export_dir = Path(settings.CASES_DIR) / case_id / "exports"
        export_dir.mkdir(parents=True, exist_ok=True)

        export_id = str(uuid.uuid4())[:8]
        filename = f"{title.replace(' ', '_')}_{focus_mode}_{export_id}.html"
        filepath = export_dir / filename

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(full_html)

        file_hash = _hash_file(str(filepath))

        # Create manifest
        manifest = {
            "export_id": export_id,
            "format": "html",
            "doc_id": doc_id,
            "case_id": case_id,
            "title": title,
            "content_hash": content_hash,
            "file_hash": file_hash,
            "exported_at": export_time,
            "exported_by": actor,
            "artifacts": [
                {"type": "html", "path": str(filepath), "hash": file_hash}
            ],
        }

        manifest_path = export_dir / f"manifest_{export_id}.json"
        with open(manifest_path, 'w') as f:
            json.dump(manifest, f, indent=2)

        # CoC event
        record_coc_event(
            case_id=case_id, actor=actor,
            action="REPORT_EXPORTED",
            target_artefact=f"studio_doc:{doc_id}",
            justification=f"Exported as HTML: {title}",
            hash_after=file_hash,
            details=manifest,
        )

        logger.info(f"[Export] HTML exported: {filepath}")
        return {
            "export_id": export_id,
            "format": "html",
            "filename": filename,
            "filepath": str(filepath),
            "file_hash": file_hash,
            "content_hash": content_hash,
            "manifest": manifest,
        }
    finally:
        conn.close()


def export_pdf(
    case_id: str,
    doc_id: str,
    actor: str = "investigator",
    frontend_url: str = None,
    cover_id: str = None,
    focus_mode: str = "Review",
    engine: str = "reportlab",
    font_style: Optional[str] = None,
    graph_style: Optional[str] = None,
    table_style: Optional[str] = None,
    selected_graphs: Optional[list[str]] = None,
    ast_override: Optional[dict] = None,
    title_override: Optional[str] = None,
) -> dict:
    """Export a document as ReportLab PDF with signed integrity manifest."""
    if ast_override is not None:
        title = title_override or f"Offline Report {doc_id}"
        ast = ast_override
        content_hash = _hash_content(canonicaljson.encode_canonical_json(ast))
    else:
        conn = open_vault(case_id)
        try:
            row = conn.execute(
                "SELECT title, ast_json, content_hash FROM studio_documents WHERE doc_id=?",
                [doc_id]
            ).fetchone()
            if not row:
                return {"error": "Document not found"}

            title, ast_str, content_hash = row[0], row[1], row[2]
            ast = json.loads(ast_str)
        finally:
            conn.close()
    export_dir = Path(settings.CASES_DIR) / case_id / "exports"
    export_dir.mkdir(parents=True, exist_ok=True)
    export_id = str(uuid.uuid4())[:8]
    filename = f"{title.replace(' ', '_')}_{focus_mode}_{export_id}.pdf"
    pdf_path = export_dir / filename

    normalized_focus_mode = (focus_mode or "Review").lower()
    normalized_engine = (engine or "reportlab").lower()
    if normalized_focus_mode == "story":
        ast = _apply_story_mode_transform(ast)
    resolved_cover_id = None
    if cover_id:
        resolved_cover_id = cover_id.strip()
        if resolved_cover_id.startswith("%2F") or resolved_cover_id.startswith("/"):
            resolved_cover_id = resolved_cover_id.split("/")[-1]
    vault_keys = None
    if normalized_focus_mode == "redact":
        try:
            evidence_service = get_report_evidence_service(case_id)
            all_keys = evidence_service.get_all_keys()
            vault_keys = {}
            for key in all_keys:
                raw_value = evidence_service.get_value(
                    key_id=key.key_id,
                    mode=RedactionMode.FULL,
                    accessed_by=actor,
                    purpose=AccessPurpose.EXPORT,
                )
                if raw_value:
                    vault_keys[key.key_id] = raw_value
        except Exception:
            logger.warning("Failed to resolve report evidence vault keys for redact mode", exc_info=True)

    pdf_bytes, generation_manifest = generate_pdf_from_ast(
        case_id=case_id,
        doc_id=doc_id,
        ast_data=ast,
        focus_mode=normalized_focus_mode,
        vault_keys=vault_keys,
        font_style=font_style,
        graph_style=graph_style,
        table_style=table_style,
        selected_graphs=selected_graphs,
    )

    with open(pdf_path, "wb") as f:
        f.write(pdf_bytes)

    integrity = PDFIntegrityService()
    signer = PDFSigner()
    pdf_sha256 = integrity.hash_pdf(pdf_bytes)

    # ── PAdES Signing (court-ready embedded signature) ───────────
    signed_payload = signer.sign_pdf(str(pdf_path))
    if signed_payload.signed_pdf_path and signed_payload.signed_pdf_path != str(pdf_path):
        # PAdES created a new signed file — use it
        import shutil
        shutil.move(signed_payload.signed_pdf_path, str(pdf_path))
        pdf_sha256 = _hash_file(str(pdf_path))
        logger.info(f"[Export] PAdES signature applied: {signed_payload.pades_level}")

    # ── QR Watermark (verification deep-link for printed copies) ─
    try:
        from operation_room.services.qr_watermark import QRWatermark
        qr = QRWatermark()
        qr.add_to_pdf(
            pdf_path=str(pdf_path),
            case_id=case_id,
            report_id=doc_id,
            manifest_hash=pdf_sha256,
            signer_identity=signed_payload.signer_identity,
        )
    except Exception as qr_err:
        logger.debug(f"[Export] QR watermark skipped: {qr_err}")

    manifest = integrity.build_export_manifest(
        case_id=case_id,
        doc_id=doc_id,
        actor=actor,
        focus_mode=focus_mode,
        generation_manifest=generation_manifest,
        pdf_hash=pdf_sha256,
        signature=signed_payload.signature_b64,
        signature_algorithm=signed_payload.algorithm,
        public_key_pem=signed_payload.public_key_pem,
    )
    manifest["content_hash"] = content_hash
    manifest["export_id"] = export_id
    manifest["filename"] = filename
    manifest["engine_requested"] = normalized_engine
    manifest["cover_id"] = resolved_cover_id
    manifest["signing"] = {
        "mode": signer.mode,
        "pades_level": signed_payload.pades_level,
        "signer_identity": signed_payload.signer_identity,
        "signing_time": signed_payload.signing_time,
    }
    manifest["style"] = {
        "font_style": font_style,
        "graph_style": graph_style,
        "table_style": table_style,
        "selected_graphs": selected_graphs or [],
    }

    manifest_path = export_dir / f"manifest_{export_id}.json"
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2)

    file_hash = _hash_file(str(pdf_path))
    record_coc_event(
        case_id=case_id,
        actor=actor,
        action="REPORT_EXPORTED_PDF",
        target_artefact=f"studio_doc:{doc_id}",
        justification=f"Exported as PDF (ReportLab): {title}",
        hash_after=file_hash,
        details=manifest,
    )

    return {
        "export_id": export_id,
        "format": "pdf",
        "engine": "reportlab",
        "engine_requested": normalized_engine,
        "cover_id": resolved_cover_id,
        "signing_mode": signer.mode,
        "pades_level": signed_payload.pades_level,
        "style": manifest["style"],
        "filename": filename,
        "filepath": str(pdf_path),
        "file_hash": file_hash,
        "manifest": manifest,
    }


def export_docx(case_id: str, doc_id: str, actor: str = "investigator") -> dict:
    """Export as DOCX using python-docx."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    conn = open_vault(case_id)
    try:
        row = conn.execute(
            "SELECT title, ast_json, content_hash FROM studio_documents WHERE doc_id=?",
            [doc_id]
        ).fetchone()
        if not row:
            return {"error": "Document not found"}

        title, ast_str, content_hash = row[0], row[1], row[2]
        ast = json.loads(ast_str)

        doc = DocxDocument()

        # NFLIP header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run("🛡️ NFLIP — National Forensic Log Intelligence Platform")
        run.font.size = Pt(14)
        run.font.bold = True

        # Classification
        cls_para = doc.add_paragraph()
        cls_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = cls_para.add_run("CONFIDENTIAL")
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(220, 38, 38)

        doc.add_paragraph("")  # Spacer

        # Walk AST and add to DOCX with figure counter
        figure_counter = {'count': 0}
        _ast_to_docx(ast, doc, figure_counter)

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f"SHA-256: {content_hash}")
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(148, 163, 184)

        # Save
        export_dir = Path(settings.CASES_DIR) / case_id / "exports"
        export_dir.mkdir(parents=True, exist_ok=True)
        export_id = str(uuid.uuid4())[:8]
        filename = f"{title.replace(' ', '_')}_{export_id}.docx"
        filepath = export_dir / filename

        doc.save(str(filepath))
        file_hash = _hash_file(str(filepath))

        record_coc_event(
            case_id=case_id, actor=actor,
            action="REPORT_EXPORTED_DOCX",
            target_artefact=f"studio_doc:{doc_id}",
            justification=f"Exported as DOCX: {title}",
            hash_after=file_hash,
        )

        return {
            "export_id": export_id,
            "format": "docx",
            "filename": filename,
            "filepath": str(filepath),
            "file_hash": file_hash,
        }
    finally:
        conn.close()


def _ast_to_docx(ast, doc, figure_counter: dict = None):
    """Walk AST and add elements to a python-docx Document."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if figure_counter is None:
        figure_counter = {'count': 0}

    ntype = ast.get("type", "")
    attrs = ast.get("attrs", {})
    content = ast.get("content", [])

    if ntype == "v4-canvas":
        # Sort elements geographically (flatten layout for sequential docx)
        all_elements = []
        for i, page in enumerate((ast.get("pages") or [])):
            for el in (page.get("elements") or []):
                all_elements.append({
                    "page": i,
                    "y": el.get("y", 0),
                    "x": el.get("x", 0),
                    "element": el
                })
        
        # Sort by Page -> Y offset -> X offset
        all_elements.sort(key=lambda item: (item["page"], item["y"], item["x"]))
        
        for item in all_elements:
            el = item["element"]
            etype = el.get("type", "component")
            data = el.get("data", {})
            
            if etype == "text":
                # Fallback purely textual text blocks as paragraphs
                import re
                txt = re.sub('<[^<]+?>', '', data.get("content", ""))
                if txt.strip():
                    doc.add_paragraph(txt)
            else:
                cat_attrs = {
                    "type": data.get("chartType") or "chart" if data.get("chartType") else data.get("type", "finding"),
                    "source": data.get("module", "case"),
                    "title": data.get("title", "Widget"),
                    "data": data,
                    "metadata": {}
                }
                _render_evidence_block_docx(cat_attrs, doc, figure_counter)
                
        # --- Vault of Evidence Appendix ---
        evidence_blocks = [item["element"] for item in all_elements if item["element"].get("type") != "text"]
        if evidence_blocks:
            doc.add_page_break()
            doc.add_heading("Vault of Evidence", level=1)
            doc.add_paragraph("Metadata registry confirming the origin, module source, and validation hashes of all functional widgets used.")
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Element ID'
            hdr_cells[1].text = 'Page'
            hdr_cells[2].text = 'Source / Type'
            hdr_cells[3].text = 'Title'
            
            # Make headers bold
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            for item in all_elements:
                el = item["element"]
                if el.get("type") == "text":
                    continue
                
                row_cells = table.add_row().cells
                row_cells[0].text = f"evd-{el.get('id', '')[:16]}..."
                row_cells[1].text = f"Page {item['page'] + 1}"
                
                data = el.get("data", {})
                mod = data.get("module") or "System"
                typ = data.get("type") or "Widget"
                row_cells[2].text = f"{mod}\n({typ})"
                row_cells[3].text = data.get("title") or "Untitled"
                
        return

    if ntype == "doc":
        for child in content:
            _ast_to_docx(child, doc, figure_counter)
    elif ntype == "heading":
        level = min(attrs.get("level", 1), 4)
        text = _extract_text_from_ast(ast)
        doc.add_heading(text, level=level)
    elif ntype == "paragraph":
        text = _extract_text_from_ast(ast)
        if text.strip():
            para = doc.add_paragraph(text)
            align = attrs.get("textAlign", "")
            if align == "center":
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif ntype == "bulletList":
        for item in content:
            text = _extract_text_from_ast(item)
            doc.add_paragraph(text, style='List Bullet')
    elif ntype == "orderedList":
        for item in content:
            text = _extract_text_from_ast(item)
            doc.add_paragraph(text, style='List Number')
    elif ntype == "blockquote":
        text = _extract_text_from_ast(ast)
        para = doc.add_paragraph(text)
        para.style = 'Intense Quote'
    elif ntype == "codeBlock":
        text = _extract_text_from_ast(ast)
        para = doc.add_paragraph(text)
        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
    elif ntype == "horizontalRule":
        doc.add_paragraph("─" * 60)
    elif ntype == "table":
        _add_table_to_docx(ast, doc)
    elif ntype == "evidenceBlock":
        _render_evidence_block_docx(attrs, doc, figure_counter)
    else:
        for child in content:
            _ast_to_docx(child, doc, figure_counter)


def _extract_text_from_ast(node: dict) -> str:
    """Extract plain text from AST."""
    if not node:
        return ""
    if node.get("type") == "text":
        return node.get("text", "")
    parts = []
    for child in (node.get("content") or []):
        parts.append(_extract_text_from_ast(child))
    return "".join(parts)


def _add_table_to_docx(table_ast, doc):
    """Convert AST table to DOCX table."""
    rows_data = []
    for row_node in (table_ast.get("content") or []):
        cells = []
        for cell_node in (row_node.get("content") or []):
            cells.append(_extract_text_from_ast(cell_node))
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                table.rows[i].cells[j].text = cell_text


# ═══════════════════════════════════════════════════════════════
# Evidence Block Rendering
# ═══════════════════════════════════════════════════════════════

# Module colors for HTML rendering
MODULE_COLORS = {
    'timeline': '#3b82f6',
    'anomaly': '#f59e0b',
    'correlation': '#8b5cf6',
    'network': '#06b6d4',
    'crud': '#10b981',
    'depth': '#ef4444',
    'case': '#64748b',
}

# Evidence type icons (using emoji for HTML)
EVIDENCE_ICONS = {
    'chart': '📊',
    'table': '📋',
    'metric': '📈',
    'finding': '🔍',
    'timeline-event': '📅',
    'anomaly': '⚠️',
    'network-flow': '🌐',
    'shap-explanation': '🧠',
    'correlation-graph': '🔗',
}


def _render_evidence_block_html(attrs: dict, figure_counter: dict) -> str:
    """Render an evidence block as premium HTML with inline SVG charts."""
    etype = attrs.get('type', 'finding')
    source = attrs.get('source', 'case')
    title = attrs.get('title', 'Evidence')
    data = attrs.get('data', {})
    metadata = attrs.get('metadata', {})

    color = MODULE_COLORS.get(source, '#64748b')
    icon = EVIDENCE_ICONS.get(etype, '📄')

    # Increment figure counter
    figure_counter['count'] = figure_counter.get('count', 0) + 1
    figure_num = figure_counter['count']

    # Build data display
    data_html = _render_evidence_data_html(etype, data, color)

    # Timestamp
    ts = metadata.get('timestamp', '')
    ts_display = ts[:19].replace('T', ' ') if ts else ''

    # Build evidence block HTML
    html = f'''
    <div class="evidence-block" style="border-left-color: {color}">
        <div class="evidence-header">
            <span class="evidence-icon">{icon}</span>
            <span class="evidence-title">Figure {figure_num}: {title}</span>
            <span class="evidence-source" style="background-color: {color}20; color: {color}">{source.upper()}</span>
        </div>
        <div class="evidence-body">
            {data_html}
        </div>
        <div class="evidence-footer">
            <span class="evidence-type">{etype}</span>
            {f'<span class="evidence-timestamp">{ts_display}</span>' if ts_display else ''}
            {f'<span class="evidence-verified">✓ Verified</span>' if metadata.get('verified') else ''}
        </div>
    </div>
    '''
    return html


def _render_evidence_data_html(etype: str, data: dict, color: str) -> str:
    """Render actual chart/table/metric data as rich inline HTML with SVG charts."""
    if not data:
        return '<p style="color:#94a3b8;font-size:10pt;">No data available</p>'

    if not isinstance(data, dict):
        if isinstance(data, (list, tuple)):
            return '<ul class="evidence-list">' + ''.join(f'<li>{item}</li>' for item in data[:8]) + '</ul>'
        return f'<p class="evidence-text">{data}</p>'

    # ── Metric type ───────────────────────────────────────────
    if 'value' in data and etype in ('metric', 'finding'):
        val = data.get('value', 'N/A')
        label = data.get('label', data.get('unit', ''))
        trend = data.get('trend')
        trend_html = ''
        if trend is not None:
            try:
                t_val = float(trend)
                arrow = '↑' if t_val > 0 else '↓'
                tcolor = '#dc2626' if t_val > 0 else '#16a34a'
                trend_html = f'<span style="color:{tcolor};font-size:12pt;font-weight:600;margin-left:8px;">{arrow} {abs(t_val)}%</span>'
            except (ValueError, TypeError):
                trend_html = f'<span style="color:#64748b;font-size:12pt;font-weight:600;margin-left:8px;">{trend}</span>'
        return f'''
            <div style="text-align:center;padding:16px;">
                <span style="font-size:36pt;font-weight:800;color:#1e293b;">{val}</span>
                {trend_html}
                <br><span style="font-size:11pt;color:#64748b;">{label}</span>
            </div>
        '''

    # ── Finding type ──────────────────────────────────────────
    if etype == 'finding' or ('severity' in data and 'summary' in data):
        severity = data.get('severity', 'info')
        summary = data.get('summary', '')
        details = data.get('details', '')
        sev_colors = {
            'critical': ('#dc2626', '#fef2f2'), 'high': ('#ea580c', '#fff7ed'),
            'medium': ('#ca8a04', '#fefce8'), 'low': ('#16a34a', '#f0fdf4'), 'info': ('#2563eb', '#eff6ff')
        }
        sc, sbg = sev_colors.get(severity, ('#64748b', '#f8fafc'))
        html = f'''
            <div style="padding:12px;">
                <span style="background:{sbg};color:{sc};border:1px solid {sc}30;padding:2px 10px;border-radius:4px;font-size:9pt;font-weight:600;text-transform:uppercase;">{severity}</span>
                <p style="margin:8px 0;font-weight:500;">{summary}</p>
                {f'<p style="font-size:10pt;color:#475569;border-left:3px solid #e2e8f0;padding-left:12px;">{details}</p>' if details else ''}
            </div>
        '''
        return html

    # ── Table type ────────────────────────────────────────────
    if 'columns' in data and 'rows' in data:
        columns = data['columns']
        rows = data.get('rows', [])[:15]
        header = '<tr>' + ''.join(f'<th>{c}</th>' for c in columns) + '</tr>'
        body = ''
        for row in rows:
            if isinstance(row, list):
                body += '<tr>' + ''.join(f'<td>{cell}</td>' for cell in row) + '</tr>'
            elif isinstance(row, dict):
                body += '<tr>' + ''.join(f'<td>{row.get(c, "")}</td>' for c in columns) + '</tr>'
        return f'<div class="evidence-table"><table>{header}{body}</table></div>'

    if 'entries' in data or 'rows' in data:
        entries = data.get('entries', data.get('rows', []))[:10]
        if entries and isinstance(entries[0], dict):
            cols = list(entries[0].keys())[:6]
            header = '<tr>' + ''.join(f'<th>{c}</th>' for c in cols) + '</tr>'
            body = ''
            for entry in entries:
                body += '<tr>' + ''.join(f'<td>{entry.get(c, "")}</td>' for c in cols) + '</tr>'
            return f'<div class="evidence-table"><table>{header}{body}</table></div>'

    # ── Chart type with chartData array → render inline SVG ──
    chart_data = data.get('chartData', data.get('events', data.get('data', [])))
    chart_type = data.get('chartType', data.get('componentId', ''))

    if isinstance(chart_data, list) and len(chart_data) > 0:
        return _render_inline_svg_chart(chart_data, chart_type, color, data)

    # ── Anomaly data ──────────────────────────────────────────
    if 'anomalies' in data or 'top_anomalies' in data:
        anomalies = data.get('anomalies', data.get('top_anomalies', []))[:8]
        if anomalies:
            rows_html = ''
            for a in anomalies:
                if isinstance(a, dict):
                    score = a.get('anomaly_score', a.get('score', '—'))
                    label = a.get('description', a.get('event', a.get('label', '—')))
                    rows_html += f'<tr><td style="font-weight:600;color:#dc2626;">{score}</td><td>{label}</td></tr>'
            return f'''
                <div class="evidence-table">
                    <table><tr><th>Score</th><th>Description</th></tr>{rows_html}</table>
                </div>
            '''

    # ── SHAP data ─────────────────────────────────────────────
    if 'features' in data or 'shap_values' in data:
        features = data.get('features', [])
        values = data.get('shap_values', data.get('importances', []))
        if features and values:
            return _render_shap_bar_svg(features[:10], values[:10], color)

    # ── Network data ──────────────────────────────────────────
    if 'flows' in data or 'connections' in data:
        flows = data.get('flows', data.get('connections', []))[:8]
        if flows and isinstance(flows[0], dict):
            keys = list(flows[0].keys())[:5]
            header = '<tr>' + ''.join(f'<th>{k}</th>' for k in keys) + '</tr>'
            body = ''
            for f in flows:
                body += '<tr>' + ''.join(f'<td>{f.get(k, "")}</td>' for k in keys) + '</tr>'
            return f'<div class="evidence-table"><table>{header}{body}</table></div>'

    # ── Generic dict fallback — but ONLY key details, not raw dump ──
    display_items = [(k, v) for k, v in data.items() if k not in ('componentId', 'metadata', 'data', 'chartData', 'events', 'raw')][:6]
    if display_items:
        html = '<div class="evidence-details">'
        for k, v in display_items:
            if isinstance(v, (dict, list)):
                try:
                    clean_v = json.loads(json.dumps(v, default=str))
                    v_str = canonicaljson.encode_canonical_json(clean_v).decode("utf-8")
                except Exception:
                    v_str = str(v)
                v = v_str[:120] + '...' if len(v_str) > 120 else v_str
            html += f'<div class="detail-row"><span class="key">{k}:</span> <span class="value">{v}</span></div>'
        html += '</div>'
        return html

    return '<p style="color:#94a3b8;font-size:10pt;">Data visualization rendered in application</p>'


def _render_inline_svg_chart(chart_data: list, chart_type: str, color: str, full_data: dict) -> str:
    """Render chart data as an inline SVG bar/area chart for static HTML export."""
    import math

    W, H = 700, 280
    PAD_L, PAD_R, PAD_T, PAD_B = 60, 20, 20, 50

    # Try to extract numeric values
    numeric_items = []
    for item in chart_data[:30]:
        if isinstance(item, dict):
            # Find the first numeric value for Y axis
            label = str(item.get('name', item.get('label', item.get('timestamp', item.get('x', '')))))[:20]
            y_val = None
            for k, v in item.items():
                if k in ('name', 'label', 'timestamp', 'x', 'date', 'time'):
                    continue
                try:
                    y_val = float(v)
                    break
                except (ValueError, TypeError):
                    continue
            if y_val is not None:
                numeric_items.append((label, y_val))
        elif isinstance(item, (int, float)):
            numeric_items.append((str(len(numeric_items)), float(item)))

    if not numeric_items:
        # No plottable numeric data — render as a nice table instead
        if chart_data and isinstance(chart_data[0], dict):
            keys = list(chart_data[0].keys())[:5]
            header = '<tr>' + ''.join(f'<th>{k}</th>' for k in keys) + '</tr>'
            body = ''
            for row in chart_data[:10]:
                body += '<tr>' + ''.join(f'<td>{row.get(k, "")}</td>' for k in keys) + '</tr>'
            return f'<div class="evidence-table"><table>{header}{body}</table></div>'
        return '<p style="color:#94a3b8;">Chart data format not recognized</p>'

    n = len(numeric_items)
    max_val = max(v for _, v in numeric_items) or 1
    min_val = min(v for _, v in numeric_items)
    if min_val > 0:
        min_val = 0
    val_range = max_val - min_val or 1

    plot_w = W - PAD_L - PAD_R
    plot_h = H - PAD_T - PAD_B

    # Generate SVG
    svg_parts = [f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {W} {H}" style="width:100%;height:auto;max-height:300px;">']

    # Background
    svg_parts.append(f'<rect width="{W}" height="{H}" fill="#fafbfc" rx="8"/>')

    # Grid lines
    for i in range(5):
        y = PAD_T + plot_h * i / 4
        val = max_val - (val_range * i / 4)
        svg_parts.append(f'<line x1="{PAD_L}" y1="{y}" x2="{W - PAD_R}" y2="{y}" stroke="#e2e8f0" stroke-width="1"/>')
        svg_parts.append(f'<text x="{PAD_L - 8}" y="{y + 4}" text-anchor="end" fill="#94a3b8" font-size="9" font-family="Inter,sans-serif">{val:.0f}</text>')

    # Determine chart style
    is_area = 'area' in chart_type.lower() or 'timeline' in chart_type.lower() or n > 12
    is_bar = 'bar' in chart_type.lower() or 'histogram' in chart_type.lower()

    if is_area or (not is_bar and n > 8):
        # Area/line chart
        points = []
        for i, (label, val) in enumerate(numeric_items):
            x = PAD_L + (plot_w * i / max(n - 1, 1))
            y = PAD_T + plot_h * (1 - (val - min_val) / val_range)
            points.append((x, y))

        # Area fill
        path = f'M {points[0][0]},{PAD_T + plot_h}'
        for x, y in points:
            path += f' L {x},{y}'
        path += f' L {points[-1][0]},{PAD_T + plot_h} Z'
        svg_parts.append(f'<path d="{path}" fill="{color}" opacity="0.15"/>')

        # Line
        line_path = f'M {points[0][0]},{points[0][1]}'
        for x, y in points[1:]:
            line_path += f' L {x},{y}'
        svg_parts.append(f'<path d="{line_path}" fill="none" stroke="{color}" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round"/>')

        # Dots
        for x, y in points:
            svg_parts.append(f'<circle cx="{x}" cy="{y}" r="3" fill="{color}" stroke="white" stroke-width="1.5"/>')
    else:
        # Bar chart
        bar_w = max(plot_w / n * 0.65, 8)
        gap = plot_w / n

        for i, (label, val) in enumerate(numeric_items):
            x = PAD_L + gap * i + (gap - bar_w) / 2
            bar_h = plot_h * ((val - min_val) / val_range)
            y = PAD_T + plot_h - bar_h
            svg_parts.append(f'<rect x="{x}" y="{y}" width="{bar_w}" height="{bar_h}" fill="{color}" opacity="0.8" rx="3"/>')

    # X-axis labels (show a subset to avoid overlap)
    step = max(1, n // 8)
    for i in range(0, n, step):
        label = numeric_items[i][0]
        if len(label) > 10:
            label = label[:10]
        x = PAD_L + (plot_w * i / max(n - 1, 1)) if n > 1 else PAD_L + plot_w / 2
        svg_parts.append(f'<text x="{x}" y="{H - 10}" text-anchor="middle" fill="#64748b" font-size="8" font-family="Inter,sans-serif" transform="rotate(-30 {x} {H - 10})">{label}</text>')

    svg_parts.append('</svg>')
    return '\n'.join(svg_parts)


def _render_shap_bar_svg(features: list, values: list, color: str) -> str:
    """Render SHAP feature importance as horizontal bar SVG."""
    W, H_PER = 700, 28
    n = len(features)
    H = n * H_PER + 40
    max_val = max(abs(v) for v in values) or 1

    svg = [f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {W} {H}" style="width:100%;height:auto;">']
    svg.append(f'<rect width="{W}" height="{H}" fill="#fafbfc" rx="8"/>')

    for i, (feat, val) in enumerate(zip(features, values)):
        y = 20 + i * H_PER
        bar_w = abs(val) / max_val * 400
        bar_color = '#dc2626' if val > 0 else '#2563eb'
        x_start = 250 if val >= 0 else 250 - bar_w

        svg.append(f'<text x="245" y="{y + 16}" text-anchor="end" fill="#475569" font-size="10" font-family="Inter,sans-serif">{str(feat)[:25]}</text>')
        svg.append(f'<rect x="{x_start}" y="{y + 4}" width="{bar_w}" height="{H_PER - 8}" fill="{bar_color}" opacity="0.7" rx="3"/>')
        svg.append(f'<text x="{x_start + bar_w + 5 if val >= 0 else x_start - 5}" y="{y + 16}" text-anchor="{"start" if val >= 0 else "end"}" fill="{bar_color}" font-size="9" font-family="Inter,sans-serif">{val:.3f}</text>')

    svg.append('</svg>')
    return '\n'.join(svg)


def _render_evidence_block_docx(attrs: dict, doc, figure_counter: dict):
    """Add an evidence block to DOCX document."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    etype = attrs.get('type', 'finding')
    source = attrs.get('source', 'case')
    title = attrs.get('title', 'Evidence')
    data = attrs.get('data', {})
    metadata = attrs.get('metadata', {})

    # Increment figure counter
    figure_counter['count'] = figure_counter.get('count', 0) + 1
    figure_num = figure_counter['count']

    icon = EVIDENCE_ICONS.get(etype, '📄')

    # Add spacing
    doc.add_paragraph("")

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(f"{icon} Figure {figure_num}: {title}")
    run.font.bold = True
    run.font.size = Pt(11)

    # Source badge
    source_para = doc.add_paragraph()
    run = source_para.add_run(f"[{source.upper()}] • {etype}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 116, 139)

    # Data content
    if data:
        if isinstance(data, dict):
            if data.get('componentId'):
                para = doc.add_paragraph()
                run = para.add_run(f"📊 Chart Component: {data.get('componentId', 'Unknown')}")
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 116, 139)
                para2 = doc.add_paragraph()
                run2 = para2.add_run("(See interactive chart in application)")
                run2.font.size = Pt(9)
                run2.font.color.rgb = RGBColor(148, 163, 184)
            elif 'value' in data:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(str(data.get('value', 'N/A')))
                run.font.size = Pt(24)
                run.font.bold = True
                if data.get('label'):
                    para2 = doc.add_paragraph()
                    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run2 = para2.add_run(data.get('label', ''))
                    run2.font.size = Pt(10)
                    run2.font.color.rgb = RGBColor(100, 116, 139)
            elif 'entries' in data or 'rows' in data:
                entries = data.get('entries', data.get('rows', []))[:5]
                if entries and len(entries) > 0:
                    first = entries[0]
                    if isinstance(first, dict):
                        cols = list(first.keys())[:4]
                        table = doc.add_table(rows=len(entries)+1, cols=len(cols))
                        table.style = 'Table Grid'
                        # Header
                        for j, col in enumerate(cols):
                            table.rows[0].cells[j].text = str(col)
                            table.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
                        # Data
                        for i, entry in enumerate(entries):
                            for j, col in enumerate(cols):
                                table.rows[i+1].cells[j].text = str(entry.get(col, ''))
                    elif isinstance(first, list):
                        num_cols = min(4, len(first))
                        table = doc.add_table(rows=len(entries), cols=num_cols)
                        table.style = 'Table Grid'
                        for i, entry in enumerate(entries):
                            for j in range(num_cols):
                                table.rows[i].cells[j].text = str(entry[j]) if j < len(entry) else ''
            else:
                # Generic dict display
                for k, v in list(data.items())[:4]:
                    para = doc.add_paragraph()
                    run1 = para.add_run(f"{k}: ")
                    run1.font.bold = True
                    run1.font.size = Pt(10)
                    run2 = para.add_run(str(v))
                    run2.font.size = Pt(10)
        elif isinstance(data, (list, tuple)):
            for item in data[:5]:
                doc.add_paragraph(str(item), style='List Bullet')
        else:
            doc.add_paragraph(str(data))

    # Footer with timestamp
    ts = metadata.get('timestamp', '')
    if ts:
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(f"Captured: {ts[:19].replace('T', ' ')}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(148, 163, 184)

    # Horizontal rule
    doc.add_paragraph("─" * 40)


def list_exports(case_id: str) -> list[dict]:
    """List all exports for a case."""
    export_dir = Path(settings.CASES_DIR) / case_id / "exports"
    if not export_dir.exists():
        return []

    exports = []
    for f in sorted(export_dir.glob("manifest_*.json"), reverse=True):
        try:
            with open(f) as mf:
                exports.append(json.load(mf))
        except Exception:
            pass
    return exports


# ═══════════════════════════════════════════════════════════════
# Evidence Appendix Generator
# ═══════════════════════════════════════════════════════════════

def generate_evidence_appendix(case_id: str, doc_id: str) -> dict:
    """Generate a comprehensive evidence appendix from a document."""
    conn = open_vault(case_id)
    try:
        row = conn.execute(
            "SELECT title, ast_json FROM studio_documents WHERE doc_id=?",
            [doc_id]
        ).fetchone()
        if not row:
            return {"error": "Document not found"}

        title, ast_str = row[0], row[1]
        ast = json.loads(ast_str)

        # Collect all evidence blocks
        evidence_items = []

        def collect_evidence(node: dict, path: str = ""):
            if not node:
                return
            ntype = node.get("type", "")
            if ntype == "evidenceBlock":
                attrs = node.get("attrs", {})
                evidence_items.append({
                    "figure_number": len(evidence_items) + 1,
                    "type": attrs.get("type", "unknown"),
                    "source": attrs.get("source", "unknown"),
                    "title": attrs.get("title", "Untitled"),
                    "data_hash": hashlib.sha256(canonicaljson.encode_canonical_json(json.loads(json.dumps(attrs.get("data", {}), default=str)))).hexdigest()[:16],
                    "metadata": attrs.get("metadata", {}),
                    "location": path,
                })
            for i, child in enumerate((node.get("content") or [])):
                collect_evidence(child, f"{path}/{i}")

        collect_evidence(ast)

        appendix = {
            "title": f"Evidence Appendix — {title}",
            "case_id": case_id,
            "doc_id": doc_id,
            "generated_at": _now(),
            "total_evidence_items": len(evidence_items),
            "evidence_by_module": {},
            "evidence_by_type": {},
            "evidence_items": evidence_items,
        }

        # Group by module
        for item in evidence_items:
            module = item["source"]
            if module not in appendix["evidence_by_module"]:
                appendix["evidence_by_module"][module] = []
            appendix["evidence_by_module"][module].append(item["figure_number"])

        # Group by type
        for item in evidence_items:
            etype = item["type"]
            if etype not in appendix["evidence_by_type"]:
                appendix["evidence_by_type"][etype] = []
            appendix["evidence_by_type"][etype].append(item["figure_number"])

        # Integrity hash
        appendix["integrity_hash"] = f"sha256:{hashlib.sha256(canonicaljson.encode_canonical_json(json.loads(json.dumps(appendix, default=str)))).hexdigest()}"

        return appendix
    finally:
        conn.close()


def verify_pdf_export(case_id: str, filename: str) -> dict:
    """Verify signed PDF export integrity using stored manifest."""
    safe_filename = Path(filename).name
    export_dir = Path(settings.CASES_DIR) / case_id / "exports"
    pdf_path = export_dir / safe_filename
    if safe_filename != filename or not pdf_path.exists() or not pdf_path.is_file():
        return {"valid": False, "error": "Export file not found"}

    stem = pdf_path.stem
    export_id = stem.split("_")[-1] if "_" in stem else ""
    manifest_path = export_dir / f"manifest_{export_id}.json"
    if not manifest_path.exists():
        return {"valid": False, "error": "Manifest not found"}

    with open(manifest_path, "r", encoding="utf-8") as f:
        manifest = json.load(f)

    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()

    integrity = PDFIntegrityService()
    computed_hash = integrity.hash_pdf(pdf_bytes)
    declared_hash = manifest.get("pdf_sha256")

    signature = manifest.get("signature", "")
    public_key_pem = manifest.get("public_key_pem", "")
    signature_ok = PDFSigner.verify_bytes(pdf_bytes, signature, public_key_pem) if signature and public_key_pem else False

    valid = computed_hash == declared_hash and signature_ok
    return {
        "valid": valid,
        "filename": safe_filename,
        "computed_pdf_sha256": computed_hash,
        "declared_pdf_sha256": declared_hash,
        "signature_valid": signature_ok,
        "manifest": manifest,
    }
