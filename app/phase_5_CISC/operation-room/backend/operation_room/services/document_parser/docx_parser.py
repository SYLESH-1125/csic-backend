"""
DOCX Document Parser for Layout Extraction.

Extracts layout information from DOCX documents:
- Paragraphs with styles
- Headers and sections
- Tables
- Images
- Lists

Used for:
- Reference document import
- Template extraction
- Layout verification
"""

import io
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from .pdf_parser import (
    DocumentLayout,
    PageLayout,
    DocumentElement,
    ElementType,
    BoundingBox,
    TextStyle,
)


class DOCXParser:
    """
    DOCX document parser.
    
    Extracts layout information from DOCX files using python-docx.
    """
    
    def __init__(self):
        """Initialize the parser."""
        self._docx_available = False
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            self._docx_available = True
        except ImportError:
            pass
    
    @property
    def is_available(self) -> bool:
        """Check if python-docx is available."""
        return self._docx_available
    
    def parse(
        self,
        source: Union[str, Path, bytes, io.BytesIO],
    ) -> DocumentLayout:
        """
        Parse a DOCX document.
        
        Args:
            source: File path, bytes, or BytesIO
            
        Returns:
            DocumentLayout with extracted elements
        """
        if not self._docx_available:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            )
        
        from docx import Document
        from docx.shared import Pt, Inches
        
        # Handle different input types
        if isinstance(source, bytes):
            source = io.BytesIO(source)
        
        doc = Document(source)
        
        layout = DocumentLayout()
        
        # Extract metadata
        if doc.core_properties:
            layout.title = doc.core_properties.title or ""
            layout.author = doc.core_properties.author or ""
            layout.metadata = {
                "created": str(doc.core_properties.created) if doc.core_properties.created else None,
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else None,
                "subject": doc.core_properties.subject,
                "keywords": doc.core_properties.keywords,
            }
        
        # Get page dimensions from section
        section = doc.sections[0] if doc.sections else None
        page_width = float(section.page_width.pt) if section else 612.0
        page_height = float(section.page_height.pt) if section else 792.0
        
        # Margins
        margin_top = float(section.top_margin.pt) if section else 72.0
        margin_bottom = float(section.bottom_margin.pt) if section else 72.0
        margin_left = float(section.left_margin.pt) if section else 72.0
        margin_right = float(section.right_margin.pt) if section else 72.0
        
        # Create initial page
        current_page = PageLayout(
            page_number=1,
            width=page_width,
            height=page_height,
            margin_top=margin_top,
            margin_bottom=margin_bottom,
            margin_left=margin_left,
            margin_right=margin_right,
        )
        
        # Track y position for element placement
        y_position = margin_top
        content_height = page_height - margin_top - margin_bottom
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            element = self._process_paragraph(paragraph, current_page.page_number)
            if element:
                # Estimate element height
                line_height = element.style.font_size * 1.2
                text_lines = max(1, len(element.content) // 80 + 1)
                element_height = line_height * text_lines
                
                # Check if we need a new page
                if y_position + element_height > page_height - margin_bottom:
                    layout.pages.append(current_page)
                    current_page = PageLayout(
                        page_number=current_page.page_number + 1,
                        width=page_width,
                        height=page_height,
                        margin_top=margin_top,
                        margin_bottom=margin_bottom,
                        margin_left=margin_left,
                        margin_right=margin_right,
                    )
                    y_position = margin_top
                
                # Set element position
                element.bbox = BoundingBox(
                    x=margin_left,
                    y=y_position,
                    width=page_width - margin_left - margin_right,
                    height=element_height,
                )
                
                current_page.elements.append(element)
                y_position += element_height + 6  # Add spacing
        
        # Process tables
        for table in doc.tables:
            table_element = self._process_table(table, current_page.page_number)
            if table_element:
                current_page.elements.append(table_element)
        
        # Add final page
        if current_page.elements:
            layout.pages.append(current_page)
        
        return layout
    
    def _process_paragraph(
        self,
        paragraph: Any,
        page_number: int,
    ) -> Optional[DocumentElement]:
        """Process a single paragraph."""
        text = paragraph.text.strip()
        if not text:
            return None
        
        # Determine style
        style = self._extract_paragraph_style(paragraph)
        
        # Determine element type
        element_type = ElementType.TEXT
        heading_level = 0
        
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        
        if "heading" in style_name or "title" in style_name:
            element_type = ElementType.HEADING
            # Extract heading level from style name
            for i in range(1, 7):
                if str(i) in style_name:
                    heading_level = i
                    break
            if heading_level == 0:
                heading_level = 1  # Default for unnamed headings
        
        elif "list" in style_name or paragraph.style and paragraph.style.name.startswith("List"):
            element_type = ElementType.LIST
        
        return DocumentElement(
            element_type=element_type,
            content=text,
            style=style,
            page_number=page_number,
            heading_level=heading_level,
        )
    
    def _extract_paragraph_style(self, paragraph: Any) -> TextStyle:
        """Extract style from paragraph."""
        from docx.shared import Pt
        
        style = TextStyle()
        
        # Get font from run
        if paragraph.runs:
            run = paragraph.runs[0]
            if run.font:
                style.font_name = run.font.name or ""
                if run.font.size:
                    style.font_size = float(run.font.size.pt)
                style.is_bold = run.font.bold or False
                style.is_italic = run.font.italic or False
        
        # Check paragraph style
        if paragraph.style:
            if paragraph.style.font:
                if paragraph.style.font.size and not style.font_size:
                    style.font_size = float(paragraph.style.font.size.pt)
                if paragraph.style.font.bold is not None:
                    style.is_bold = paragraph.style.font.bold
                if paragraph.style.font.italic is not None:
                    style.is_italic = paragraph.style.font.italic
        
        return style
    
    def _process_table(
        self,
        table: Any,
        page_number: int,
    ) -> DocumentElement:
        """Process a table."""
        # Extract table content
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            rows.append(cells)
        
        return DocumentElement(
            element_type=ElementType.TABLE,
            content=f"Table with {len(rows)} rows",
            page_number=page_number,
            metadata={
                "rows": rows,
                "row_count": len(rows),
                "col_count": len(rows[0]) if rows else 0,
            },
        )
    
    def extract_text(
        self,
        source: Union[str, Path, bytes, io.BytesIO],
    ) -> str:
        """
        Extract plain text from DOCX.
        
        Args:
            source: File path, bytes, or BytesIO
            
        Returns:
            Extracted text
        """
        if not self._docx_available:
            raise ImportError("python-docx is required")
        
        from docx import Document
        
        if isinstance(source, bytes):
            source = io.BytesIO(source)
        
        doc = Document(source)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        return "\n".join(text_parts)
    
    def get_structure(
        self,
        source: Union[str, Path, bytes, io.BytesIO],
    ) -> List[Dict[str, Any]]:
        """
        Get document structure (headings and sections).
        
        Args:
            source: File path, bytes, or BytesIO
            
        Returns:
            List of structure elements
        """
        if not self._docx_available:
            raise ImportError("python-docx is required")
        
        from docx import Document
        
        if isinstance(source, bytes):
            source = io.BytesIO(source)
        
        doc = Document(source)
        
        structure = []
        for i, paragraph in enumerate(doc.paragraphs):
            style_name = paragraph.style.name if paragraph.style else ""
            
            if "heading" in style_name.lower() or "title" in style_name.lower():
                level = 0
                for j in range(1, 7):
                    if str(j) in style_name:
                        level = j
                        break
                
                structure.append({
                    "type": "heading",
                    "text": paragraph.text,
                    "level": level or 1,
                    "index": i,
                })
        
        return structure
